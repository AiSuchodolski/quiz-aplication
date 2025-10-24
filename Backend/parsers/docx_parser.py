from pathlib import Path
import docx


def docx_parser():
    docx_path = list(Path("..", "upload").glob("*.docx"))
    if not docx_path:
        print("No DOCX files found in uploads folder")
        exit()
    data_path = Path("..", "data", "data.txt")
    data_path.parent.mkdir(exist_ok=True)
    with data_path.open("a", encoding="utf-8", errors="ignore") as open_data_file:
        for docx_file in docx_path:
            doc = docx.Document(docx_file)
            for paragraph in doc.paragraphs:
                open_data_file.write(f'{paragraph.text}\n')
    print(f"DOCX files parsed successfully and saved to {data_path}")


docx_parser()












#     docx_path = os.path.join("..", "upload", "*.docx")
#     docx_files = glob.glob(docx_path)
#     if not docx_files:
#         print(f"No DOCX file found in {docx_path}")
#         exit()
#     data_path = os.path.join("..", "data", "data.txt")
#     for docx_file in docx_files:
#         doc = docx.Document(docx_file)
#         with open(data_path, "a", encoding="utf-8") as data_file:
#             for paragraph in doc.paragraphs:
#                data_file.write(f'{paragraph.text}\n')

    
#     print("DOCX file parsed successfully and saved to data.txt")

# docx_parser()